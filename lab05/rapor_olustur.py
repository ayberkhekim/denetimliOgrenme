# -*- coding: utf-8 -*-
"""
Lab05 Rapor Olusturucu - DOCX formatinda profesyonel rapor olusturur.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# Sayfa kenarliklari
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

OUTPUT = "outputs"

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_image_safe(path, width=Inches(5.5)):
    full = os.path.join(OUTPUT, path)
    if os.path.exists(full):
        doc.add_picture(full, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Gorsel bulunamadi: {path}]")

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ============================================================
# KAPAK SAYFASI
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("YZM0206 - Yapay Zeka ve Makine Ogrenmesi")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Laboratuvar Foyusu 5")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dijital Goruntu Isleme")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x53, 0x34, 0x83)

for _ in range(4):
    doc.add_paragraph()

info_data = [
    ("Ad Soyad", "___________________"),
    ("Ogrenci No", "___________________"),
    ("Tarih", "30.04.2026"),
]
for label, value in info_data:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(12)
    p.add_run(value).font.size = Pt(12)

doc.add_page_break()

# ============================================================
# ICINDEKILER
# ============================================================
add_heading_styled("Icindekiler", level=1)
toc_items = [
    "1. Amac ve Temel Kavramlar",
    "2. Kullanilan Kutuphaneler",
    "3. Uygulama 1: Noktasal Islemler ve Renk Uzaylari",
    "4. Uygulama 2: Uzamsal Filtreleme (Convolution)",
    "5. Uygulama 3: Otsu Esikleme Metodu",
    "6. Uygulama 4: RGB ve Grayscale Iliskisi",
    "7. Bonus Egzersizler",
    "8. Sonuc ve Degerlendirme",
    "9. Referanslar",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 1. AMAC
# ============================================================
add_heading_styled("1. Amac ve Temel Kavramlar", level=1)
doc.add_paragraph(
    "Bu laboratuvar calismasinin amaci, dijital goruntuleme prensiplerini anlamak, "
    "uzamsal filtreleme (convolution) tekniklerini uygulamak ve goruntu "
    "butleme yontemlerini analiz etmektir."
)
doc.add_paragraph(
    "Dijital bir goruntu, f(x,y) seklinde tanimlanabilen iki boyutlu ayrik bir fonksiyondur. "
    "Burada x ve y uzamsal koordinatlar, f degeri ise o noktadaki yogunluk (intensity) "
    "degerini temsil eder. Renkli goruntulerde her piksel 3 kanalli (RGB) bir vektordur."
)

# ============================================================
# 2. KUTUPHANELER
# ============================================================
add_heading_styled("2. Kullanilan Kutuphaneler", level=1)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Kutuphane", "Surum", "Kullanim Amaci"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ("OpenCV (cv2)", "4.13.0", "Goruntu okuma, donusturme, filtreleme"),
    ("NumPy", "2.4.4", "Matris islemleri, sayisal hesaplama"),
    ("Matplotlib", "3.10.9", "Gorsellestirme, grafik cizimi"),
]
for i, (lib, ver, usage) in enumerate(data):
    table.rows[i+1].cells[0].text = lib
    table.rows[i+1].cells[1].text = ver
    table.rows[i+1].cells[2].text = usage

doc.add_paragraph()
add_code_block("import cv2 as cv\nimport numpy as np\nimport matplotlib.pyplot as plt")

doc.add_page_break()

# ============================================================
# 3. UYGULAMA 1
# ============================================================
add_heading_styled("3. Uygulama 1: Noktasal Islemler ve Renk Uzaylari", level=1)

add_heading_styled("3.1 Renk Uzayi Donusumleri", level=2)
doc.add_paragraph(
    "Lena goruntusu OpenCV ile okunarak RGB, BGR, YCrCb ve HSV renk uzaylarina "
    "donusturulmustur. OpenCV varsayilan olarak BGR formatinda okur; matplotlib ile "
    "gostermek icin RGB'ye cevirmek gerekir."
)
doc.add_paragraph(
    "YCrCb: Y parlakligi (luminance), Cr kirmizi-fark, Cb mavi-fark bilesenlerini icerir. "
    "HSV: H renk tonu (hue), S doygunluk (saturation), V deger (brightness) bilesenlerinden olusur."
)

add_code_block(
    "img_bgr = cv.imread('lena.jpg')\n"
    "img_rgb = cv.cvtColor(img_bgr, cv.COLOR_BGR2RGB)\n"
    "img_ycrcb = cv.cvtColor(img_bgr, cv.COLOR_BGR2YCrCb)\n"
    "img_hsv = cv.cvtColor(img_bgr, cv.COLOR_BGR2HSV)"
)
add_image_safe("renk_uzaylari.png", Inches(5.8))

add_heading_styled("3.2 Negatif Alma", level=2)
doc.add_paragraph(
    "Negatif alma islemi: g(x,y) = 255 - f(x,y). Bu, lineer cebirde bir afin "
    "donusumdur: G = 255*J - F (J: tum elemanlari 1 olan matris). "
    "Her pikselin parlaklik degeri ters cevrilerek goruntunun negatifi elde edilir."
)
add_code_block("img_negative = 255 - img_rgb  # NumPy broadcasting")
add_image_safe("negatif.png", Inches(5.5))

add_heading_styled("3.3 Histogram ve Istatistiksel Analiz", level=2)
doc.add_paragraph(
    "Goruntunun histogrami, piksel degerlerinin dagilimini gosterir. "
    "Asagida hem grayscale hem de RGB kanal histogramlari verilmistir."
)
add_image_safe("histogram.png", Inches(5.5))

doc.add_paragraph()
stat_table = doc.add_table(rows=5, cols=5)
stat_table.style = 'Light Grid Accent 1'
stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

stat_headers = ["Kanal", "Min", "Max", "Median", "Mean"]
for i, h in enumerate(stat_headers):
    stat_table.rows[0].cells[i].text = h
    for p in stat_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

stat_data = [
    ("Red (R)", "0", "255", "177.00", "154.24"),
    ("Green (G)", "0", "255", "110.00", "111.51"),
    ("Blue (B)", "0", "255", "85.00", "93.46"),
    ("Grayscale", "0", "255", "129.00", "122.24"),
]
for i, row_data in enumerate(stat_data):
    for j, val in enumerate(row_data):
        stat_table.rows[i+1].cells[j].text = val

doc.add_page_break()

# ============================================================
# 4. UYGULAMA 2
# ============================================================
add_heading_styled("4. Uygulama 2: Uzamsal Filtreleme (Convolution)", level=1)

doc.add_paragraph(
    "Convolution (gezdirme) islemi, goruntu matrisinin her piksel komsulugu ile "
    "bir kernel (filtre) matrisinin eleman-eleman carpilip toplanmasidir: "
    "g(x,y) = SUM_i SUM_j h(i,j) * f(x-i, y-j). "
    "Kenar piksellerinde veri kaybini onlemek icin Zero Padding kullanilir."
)

add_heading_styled("4.1 Gaussian ve Mean Filtresi Karsilastirmasi", level=2)
doc.add_paragraph(
    "Gaussian kernel, merkeze daha fazla agirlik veren agirlikli bir ortalamadir: "
    "h(x,y) = (1/2*pi*sigma^2) * exp(-(x^2+y^2)/(2*sigma^2)). "
    "Mean filtresi ise tum komsulara esit agirlik verir."
)
add_code_block(
    "gaussian_3x3 = cv.GaussianBlur(img_gray, (3, 3), 0)\n"
    "gaussian_5x5 = cv.GaussianBlur(img_gray, (5, 5), 0)\n"
    "mean_3x3 = cv.blur(img_gray, (3, 3))\n"
    "mean_5x5 = cv.blur(img_gray, (5, 5))"
)
add_image_safe("gaussian_vs_mean.png", Inches(5.5))

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Teknik Yorum - 3x3 vs 5x5 Filtre Boyutu:")
run.bold = True
doc.add_paragraph(
    "Filtre boyutu arttikca yumusatma etkisi artar ve goruntu daha bulanik olur. "
    "Ancak kenar detaylari (edge details) kaybolur. Hesaplama maliyeti de artar: "
    "3x3 = 9 carpma islemi, 5x5 = 25 carpma islemi. "
    "Gaussian filtresi merkeze yakin piksellere daha fazla agirlik verdigi icin "
    "daha dogal bir yumusatma saglarken, Mean filtresi tum piksellere esit agirlik "
    "verip daha agresif bir bulaniklastirma yapar."
)

add_heading_styled("4.2 Laplacian Filtresi (Kenar Tespiti)", level=2)
doc.add_paragraph(
    "Laplacian, ikinci turev operatorudur. Yuksek frekans bilesenlerini (kenarlar) tespit eder. "
    "Ayrik formda kernel: [[0,1,0],[1,-4,1],[0,1,0]]."
)
add_code_block("laplacian = cv.Laplacian(img_gray, cv.CV_64F)")
add_image_safe("laplacian_filtre.png", Inches(5.5))

add_heading_styled("4.3 Median Filtresi (Gurultu Temizleme)", level=2)
doc.add_paragraph(
    "Median filtre, lineer olmayan bir filtredir. Komsuluk icindeki piksel degerlerini "
    "siralar ve ortanca degeri secer. Salt-and-pepper (tuz-biber) gurultusune karsi "
    "cok etkilidir ve kenarlari koruyarak gurultuyu temizler."
)
add_code_block(
    "# Gurultu ekleme\n"
    "img_noisy = add_salt_pepper_noise(img_gray, amount=0.05)\n"
    "# Median filtre uygulama\n"
    "median_3x3 = cv.medianBlur(img_noisy, 3)\n"
    "median_5x5 = cv.medianBlur(img_noisy, 5)"
)
add_image_safe("median_filtre.png", Inches(5.5))

add_heading_styled("4.4 Padding (Sifir Ekleme)", level=2)
doc.add_paragraph(
    "Kenar piksellerinde filtre gezdirirken veri kaybini onlemek icin goruntunun "
    "cevresine sifir (0) degerli pikseller eklenir. Padding miktari: p = (kernel_size - 1) / 2. "
    "3x3 kernel icin 1 piksel, 5x5 kernel icin 2 piksel padding uygulanir."
)
add_image_safe("padding.png", Inches(5.0))

doc.add_page_break()

# ============================================================
# 5. UYGULAMA 3
# ============================================================
add_heading_styled("5. Uygulama 3: Otsu Esikleme Metodu", level=1)
doc.add_paragraph(
    "Otsu metodu, histogrami iki sinifa ayiran optimal esik degerini bulur. "
    "Sinif-ici varyansi minimize eden (veya esdeger olarak sinif-arasi varyansi "
    "maximize eden) esik degerini secer. Bu tam bir arama (exhaustive search) olup "
    "t = [0, 255] araliginda en iyi esik degerini bulur."
)
add_code_block(
    "img_gray = cv.cvtColor(img_bgr, cv.COLOR_BGR2GRAY)\n"
    "otsu_thresh, img_binary = cv.threshold(\n"
    "    img_gray, 0, 255, cv.THRESH_BINARY + cv.THRESH_OTSU)"
)
add_image_safe("otsu_esikleme.png", Inches(5.5))

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Sonuclar:")
run.bold = True
doc.add_paragraph("Otsu tarafindan bulunan esik degeri: 112")
doc.add_paragraph("Beyaz piksel orani: %58.9 | Siyah piksel orani: %41.1")

doc.add_paragraph()
doc.add_paragraph(
    "Asagida farkli sabit esik degerleri (100, 150, 200) ile Otsu'nun otomatik "
    "buldugu deger karsilastirilmistir:"
)
add_image_safe("esik_karsilastirma.png", Inches(5.5))

doc.add_page_break()

# ============================================================
# 6. UYGULAMA 4
# ============================================================
add_heading_styled("6. Uygulama 4: RGB ve Grayscale Iliskisi", level=1)

add_heading_styled("6.1 Kanal Ayirimi", level=2)
doc.add_paragraph(
    "RGB goruntunun R, G ve B kanallari ayri ayri cikarilmis ve gorsellestrilmistir. "
    "Her kanal hem grayscale hem de kendi renginde gosterilmistir."
)
add_code_block(
    "R = img_rgb[:, :, 0]  # Red kanali\n"
    "G = img_rgb[:, :, 1]  # Green kanali\n"
    "B = img_rgb[:, :, 2]  # Blue kanali"
)
add_image_safe("kanal_ayrimi.png", Inches(5.8))

add_heading_styled("6.2 Grayscale Donusum Formulu", level=2)
doc.add_paragraph(
    "Grayscale donusum formulu: Gray = 0.299*R + 0.587*G + 0.114*B. "
    "Katsayilar insan gorsel algi sistemine (HVS) dayanir. Yesil kanali en yuksek "
    "agirliga sahiptir (%58.7) cunku insan gozu yesile en duyarlidir."
)
add_image_safe("grayscale_dogrulama.png", Inches(5.5))

doc.add_paragraph(
    "Manuel hesaplama ile OpenCV arasindaki maksimum fark 1 piksel, "
    "ortalama fark 0.5122'dir. Bu minimal fark, yuvarlama hassasiyetinden kaynaklanir."
)

add_heading_styled("6.3 Kanallarin Grayscale'e Katkisi", level=2)
doc.add_paragraph(
    "Asagida her kanalin agirlikli katkisi gosterilmistir. "
    "Yesil kanalin en parlak (en yuksek katki) gorundugu acikca gorulmektedir."
)
add_image_safe("kanal_katkilari.png", Inches(5.5))

doc.add_page_break()

# ============================================================
# 7. BONUS
# ============================================================
add_heading_styled("7. Bonus Egzersizler", level=1)
doc.add_paragraph(
    "Referans sitesindeki (Universidad de Oviedo) 7 egzersiz tamamlanmistir."
)

exercises = [
    ("7.1 Exercise 1: Bolge Cikarma", 
     "Goruntuden belirli bir bolgeyi cikarip float32 formatinda donduren fonksiyon yazilmistir. "
     "cameraman.tif uzerinde kameracinin basi cikarilmistir.",
     "exercise1_bolge_cikarma.png"),
    ("7.2 Exercise 2: Dairesel Maske",
     "lena_gray_512.tif uzerinde yaricapi 150 piksel olan dairesel bir maske uygulanmistir. "
     "Daire disi pikseller hem sifirlanmis hem de yarim yogunlukla gosterilmistir. "
     "Maske, (j-cx)^2 + (i-cy)^2 < r^2 kosulu ile olusturulmustur.",
     "exercise2_dairesel_maske.png"),
    ("7.3 Exercise 3: Lineer Degradasyon",
     "np.linspace(1, 0, rows) ile 1'den 0'a azalan vektor olusturulup "
     "np.tile ile matrise genisletilerek Hadamard carpimi uygulanmistir.",
     "exercise3_degradasyon.png"),
    ("7.4 Exercise 4: Satranc Tahtasi",
     "250x250 piksellik karelerden olusan 8x8 satranc tahtasi np.tile ve np.kron ile olusturulmustur.",
     "exercise4_satranc.png"),
    ("7.5 Exercise 5: Es Merkezli Daireler",
     "500x500 piksellik goruntu uzerinde ~10 piksel genisliginde es merkezli daireler "
     "olusturulmustur. Merkeze uzaklik ve modulo islemi kullanilmistir.",
     "exercise5_daireler.png"),
    ("7.6 Exercise 6: Pecete Deseni",
     "10x10 piksellik karelerden olusan pecete deseni np.tile ile olusturulmustur.",
     "exercise6_pecete.png"),
    ("7.7 Exercise 7: Noktalar Deseni",
     "500x500 goruntu uzerinde yaricapi 10, merkezler arasi 50 piksel olan "
     "duzgun aralikli daireler np.tile ile olusturulmustur.",
     "exercise7_noktalar.png"),
]

for title, desc, img_name in exercises:
    add_heading_styled(title, level=2)
    doc.add_paragraph(desc)
    add_image_safe(img_name, Inches(4.5))
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 8. SONUC
# ============================================================
add_heading_styled("8. Sonuc ve Degerlendirme", level=1)
doc.add_paragraph(
    "Bu laboratuvar calismasinda dijital goruntu islemenin temel kavramlari "
    "uygulamali olarak incelenmistir. Elde edilen bulgular:"
)

conclusions = [
    "Renk Uzaylari: Farkli renk uzaylari (RGB, YCrCb, HSV) farkli uygulamalarda avantaj saglar. "
    "Ornegin HSV, renk tabanli segmentasyon icin daha uygundur.",
    "Uzamsal Filtreleme: Filtre boyutu arttikca yumusatma artar ancak detay kaybi da artar. "
    "Gaussian filtresi, Mean filtresine gore daha dogal sonuclar uretir.",
    "Otsu Esikleme: Otomatik esik belirleme icin etkili bir yontemdir. "
    "Lena goruntusu icin optimal esik 112 olarak bulunmustur.",
    "RGB-Grayscale: Grayscale donusum, kanallarin agirlikli ortalamasi ile elde edilir. "
    "Yesil kanalin en yuksek katkiya sahip olmasi, insan gorsel algisiyla uyumludur.",
    "Bonus Egzersizler: np.tile ve np.kron gibi vektorize islemler, dongu tabanli "
    "yaklasimlara gore cok daha hizli ve okunakli kod uretmektedir.",
]
for c in conclusions:
    doc.add_paragraph(c, style='List Bullet')

# ============================================================
# 9. REFERANSLAR
# ============================================================
add_heading_styled("9. Referanslar", level=1)
refs = [
    "OpenCV Documentation - https://docs.opencv.org/",
    "Escuela de Ingenieria Informatica / Universidad de Oviedo - "
    "Intro to Image Processing: https://www.unioviedo.es/compnum/labs/new/intro_image.html",
    "Gonzalez, R. C. & Woods, R. E. - Digital Image Processing (4th Edition)",
    "NumPy Documentation - https://numpy.org/doc/",
    "Matplotlib Documentation - https://matplotlib.org/stable/",
]
for ref in refs:
    doc.add_paragraph(ref, style='List Number')

# ============================================================
# KAYDET
# ============================================================
output_path = "Lab05_Rapor.docx"
doc.save(output_path)
print(f"Rapor basariyla olusturuldu: {os.path.abspath(output_path)}")
